"""Hide a message in another text using 'invisible ink'."""
import docx
from docx.shared import RGBColor, Pt

fake_text = docx.Document('fakeMessage.docx')
fake_list = []
for paragraph in fake_text.paragraphs:
    fake_list.append(paragraph.text)

real_text = docx.Document('realMessage.docx')
real_list = []
for paragraph in real_text.paragraphs:
    if len(paragraph.text) != 0:       # remove blank lines
        real_list.append(paragraph.text)

doc = docx.Document('template.docx')

doc.add_heading('Morland Holmes', 0)
subtitle = doc.add_heading('Global Consulting & Negotiations', 1)
subtitle.alignment = 1
doc.add_heading('', 1)
doc.add_paragraph('December 17, 2015')
doc.add_paragraph('')

def set_spacing(pgraph):
    """Use docx to set line spacing between paragraphs."""
    paragraph_format = pgraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

LENGTH_REAL = len(real_list)
COUNT_REAL = 0    # index of current line in real (hidden) message

for line in fake_list:
    if COUNT_REAL < LENGTH_REAL and line == "":
        paragraph = doc.add_paragraph(real_list[COUNT_REAL])
        paragraph_index = len(doc.paragraphs) - 1
        run = doc.paragraphs[paragraph_index].runs[0]
        font = run.font
        font.color.rgb = RGBColor(255,255,255)
        COUNT_REAL += 1
    else:
        paragraph = doc.add_paragraph(line)

    set_spacing(paragraph)

doc.save('ciphertext_message_letterhead.docx')

print("Done")
